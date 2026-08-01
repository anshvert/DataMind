from pathlib import Path

from docx import Document as DocxDocument
import pandas as pd

from naturallangdata.core.config import Settings
from naturallangdata.services.pdf_extractor import PDFExtractionService
from naturallangdata.services.tabular_store import TabularStore


class DocumentExtractionService:
    """Extract text from supported document types (PDF, DOCX, CSV, XLSX)."""

    def __init__(self, pdf_extractor: PDFExtractionService, settings: Settings) -> None:
        self._pdf_extractor = pdf_extractor
        self._settings = settings
        self._tabular_store = TabularStore(settings.tabular_db_path)

    def extract(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._pdf_extractor.extract(file_path)
        if suffix == ".docx":
            return self._extract_docx(file_path)
        if suffix == ".csv":
            return self._extract_csv(file_path)
        if suffix == ".xlsx":
            return self._extract_xlsx(file_path)
        raise ValueError(f"Unsupported file type: {suffix}")

    def delete_document(self, doc_id: str) -> None:
        self._tabular_store.delete_document(doc_id)

    def _extract_docx(self, file_path: Path) -> str:
        doc = DocxDocument(str(file_path))

        lines: list[str] = []
        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if text:
                lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                row_text = " | ".join(part for part in cells if part)
                if row_text:
                    lines.append(row_text)

        return "\n".join(lines).strip()

    def _extract_csv(self, file_path: Path) -> str:
        frame = pd.read_csv(file_path)
        return self._build_tabular_text(file_path=file_path, sheets={"sheet1": frame})

    def _extract_xlsx(self, file_path: Path) -> str:
        sheets = pd.read_excel(file_path, sheet_name=None)
        if not isinstance(sheets, dict) or not sheets:
            return ""
        return self._build_tabular_text(file_path=file_path, sheets=sheets)

    def _build_tabular_text(self, file_path: Path, sheets: dict[str, pd.DataFrame]) -> str:
        doc_id = file_path.stem
        sections: list[str] = []

        for sheet_name, frame in sheets.items():
            safe_sheet_name = sheet_name or "sheet"
            table_name = self._tabular_store.ingest_dataframe(
                doc_id=doc_id,
                sheet_name=safe_sheet_name,
                frame=frame,
            )

            rows, cols = frame.shape
            sections.append(
                f"[TABLE {safe_sheet_name}] rows={rows} columns={cols} sqlite_table={table_name}."
            )
            sections.append(self._schema_summary(safe_sheet_name, frame))
            sections.extend(self._row_chunks(safe_sheet_name, frame))

        return "\n\n".join(part for part in sections if part).strip()

    def _schema_summary(self, sheet_name: str, frame: pd.DataFrame) -> str:
        column_defs = [f"{col} ({dtype})" for col, dtype in frame.dtypes.items()]
        return f"Schema for {sheet_name}: " + ", ".join(column_defs)

    def _row_chunks(self, sheet_name: str, frame: pd.DataFrame) -> list[str]:
        if frame.empty:
            return [f"No rows found in table {sheet_name}."]

        max_rows = min(len(frame), self._settings.tabular_max_rows_for_embedding)
        batch_size = max(1, self._settings.tabular_row_batch_size)
        max_cols = max(1, self._settings.tabular_max_columns_per_row)
        max_cell_chars = max(20, self._settings.tabular_max_cell_chars)

        working = frame.head(max_rows).fillna("")
        selected_columns = [str(col) for col in working.columns[:max_cols]]
        row_chunks: list[str] = []

        for start in range(0, len(working), batch_size):
            stop = min(start + batch_size, len(working))
            lines = [f"Table {sheet_name} rows {start + 1} to {stop}."]

            for idx in range(start, stop):
                record = working.iloc[idx]
                parts: list[str] = []
                for col in selected_columns:
                    value = str(record[col]).strip().replace("\n", " ")
                    if len(value) > max_cell_chars:
                        value = value[:max_cell_chars] + "..."
                    parts.append(f"{col}={value}")

                lines.append(f"Row {idx + 1}. " + "; ".join(parts) + ".")

            if len(working.columns) > max_cols:
                lines.append(
                    f"Columns truncated to first {max_cols} out of {len(working.columns)} for embedding efficiency."
                )

            row_chunks.append("\n".join(lines))

        if len(frame) > max_rows:
            row_chunks.append(
                f"Only first {max_rows} rows embedded out of {len(frame)} total rows for efficiency."
            )

        return row_chunks
